from typing import List, Dict, Optional, BinaryIO
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListItem, ListFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
import markdown
import re
from ..config import Config, OutputFormat

class FormatterService:
    def __init__(self):
        self.pdf_font_size = Config.PDF_FONT_SIZE
        self.docx_template = Config.DOCX_TEMPLATE

    def format_output(self, 
                     content: str, 
                     sources: List[str], 
                     output_format: OutputFormat,
                     title: Optional[str] = None) -> Dict[str, any]:
        """
        Format the research report in the specified format
        """
        if output_format not in Config.VALID_OUTPUT_FORMATS:
            raise ValueError(f"Invalid output format: {output_format}")

        if output_format == "text":
            return self._format_text(content, sources, title)
        elif output_format == "markdown":
            return self._format_markdown(content, sources, title)
        elif output_format == "pdf":
            return self._format_pdf(content, sources, title)
        elif output_format == "docx":
            return self._format_docx(content, sources, title)

    def _format_text(self, content: str, sources: List[str], title: Optional[str]) -> Dict[str, any]:
        """Format as plain text"""
        formatted_content = []
        
        if title:
            formatted_content.append(title)
            formatted_content.append("=" * len(title))
            formatted_content.append("")
        
        formatted_content.append(content)
        formatted_content.append("\nSources:")
        formatted_content.extend([f"- {source}" for source in sources])
        
        return {
            "content": "\n".join(formatted_content),
            "format": "text",
            "mime_type": "text/plain"
        }

    def _format_markdown(self, content: str, sources: List[str], title: Optional[str]) -> Dict[str, any]:
        """Format as markdown"""
        formatted_content = []
        
        if title:
            formatted_content.append(f"# {title}")
            formatted_content.append("")
        
        formatted_content.append(content)
        formatted_content.append("\n## Sources")
        formatted_content.extend([f"* {source}" for source in sources])
        
        return {
            "content": "\n".join(formatted_content),
            "format": "markdown",
            "mime_type": "text/markdown"
        }

    def _convert_markdown_to_html(self, text: str) -> str:
        """Convert markdown to HTML with extensions"""
        return markdown.markdown(text, extensions=['extra'])

    def _clean_html_for_pdf(self, html: str) -> str:
        """Clean HTML for PDF rendering"""
        # Remove multiple newlines
        html = re.sub(r'\n\s*\n', '\n\n', html)
        # Convert HTML lists to plain text bullets
        html = re.sub(r'<li>', '• ', html)
        html = re.sub(r'</li>', '\n', html)
        # Remove other HTML tags
        html = re.sub(r'<[^>]+>', '', html)
        return html.strip()

    def _format_pdf(self, content: str, sources: List[str], title: Optional[str]) -> Dict[str, any]:
        """Format as PDF with proper markdown rendering"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Create styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=16
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=self.pdf_font_size,
            leading=self.pdf_font_size * 1.2,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        ))
        
        story = []
        
        # Add title
        if title:
            story.append(Paragraph(title, styles['CustomTitle']))
        
        # Convert markdown to HTML and clean it
        html_content = self._convert_markdown_to_html(content)
        cleaned_content = self._clean_html_for_pdf(html_content)
        
        # Process content by sections (split by markdown headers)
        sections = re.split(r'(#{1,6}\s.*)', cleaned_content)
        for section in sections:
            if section.strip():
                if section.startswith('#'):
                    # Count heading level and clean the header text
                    level = len(re.match(r'#+', section).group())
                    header_text = section.lstrip('#').strip()
                    if level == 1:
                        story.append(Paragraph(header_text, styles['CustomHeading1']))
                    else:
                        story.append(Paragraph(header_text, styles['CustomHeading2']))
                else:
                    # Process paragraphs
                    paragraphs = section.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            # Handle bullet points
                            if paragraph.strip().startswith('•'):
                                items = [ListItem(Paragraph(item.strip('• '), styles['CustomBody']))
                                       for item in paragraph.split('\n') if item.strip()]
                                story.append(ListFlowable(items, bulletType='bullet', leftIndent=20))
                            else:
                                story.append(Paragraph(paragraph, styles['CustomBody']))
        
        # Add sources section
        story.append(Spacer(1, 20))
        story.append(Paragraph('Sources', styles['CustomHeading2']))
        for source in sources:
            story.append(Paragraph(f"• {source}", styles['CustomBody']))
        
        # Build PDF
        doc.build(story)
        
        return {
            "content": buffer.getvalue(),
            "format": "pdf",
            "mime_type": "application/pdf"
        }

    def _format_docx(self, content: str, sources: List[str], title: Optional[str]) -> Dict[str, any]:
        """Format as DOCX"""
        if self.docx_template and os.path.exists(self.docx_template):
            doc = Document(self.docx_template)
        else:
            doc = Document()

        if title:
            doc.add_heading(title, 0)

        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        # Add sources
        doc.add_heading('Sources', level=2)
        for source in sources:
            doc.add_paragraph(f"• {source}", style='List Bullet')

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        
        return {
            "content": buffer.getvalue(),
            "format": "docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        } 